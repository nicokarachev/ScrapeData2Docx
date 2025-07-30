# https://research.com/journal/the-lancet
# https://research.com/journal/new-england-journal-of-medicine
# https://research.com/journal/nature-medicine
# https://research.com/journal/the-lancet-oncology
# https://research.com/journal/jama-journal-of-the-american-medical-association

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
import time

options = Options()
options.binary_location = "/usr/bin/google-chrome"
options.add_argument("--headless=new")
options.add_argument("--no-sandbox")
options.add_argument("--disable-dev-shm-usage")
options.add_argument("--disable-gpu")
options.add_argument("--disable-extensions")
options.add_argument("user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115 Safari/537.36")

service = Service(ChromeDriverManager().install())
driver = webdriver.Chrome(service=service, options=options)

# Go to the target URL
url = "https://research.com/journal/the-lancet"
print(url)
driver.get(url)

try:
    WebDriverWait(driver, 30).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
except:
    print("Timeout loading page")
# Get page content
html = driver.page_source
# print(html)
time.sleep(3)

soup = BeautifulSoup(html, "html.parser")
# print(soup)

# Example: Extract all links
desc = soup.find("div", class_="tab-slide")
# print(desc)

topics = desc.find("h2")

print(topics)

if desc:
    print("Journal Summary:", topics.getText())
else:
    print("Summary not found.")

# Clean up
driver.quit()

# Create a new Word document
doc = Document()

# Add title and paragraphs
doc.add_heading('My Report', level=1)
doc.add_paragraph('This is the first paragraph.')
doc.add_paragraph('Another paragraph with some more content.')
doc.add_paragraph(topics)

# Save the document
doc.save('my_report.docx')
